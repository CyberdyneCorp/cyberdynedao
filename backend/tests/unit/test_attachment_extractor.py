"""MultiFormatTextExtractor: real extraction per supported format."""

from __future__ import annotations

import csv
import io

import pytest
from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas

from cyberdyne_backend.adapters.outbound.attachments import MultiFormatTextExtractor

pytestmark = pytest.mark.asyncio

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


async def test_extract_csv_renders_rows() -> None:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["name", "score"])
    writer.writerow(["ada", "99"])
    data = buf.getvalue().encode("utf-8")

    text = await MultiFormatTextExtractor().extract(data=data, content_type="text/csv")

    assert "name" in text
    assert "ada" in text
    assert "99" in text


async def test_extract_xlsx_renders_cells() -> None:
    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Sheet1"
    ws.append(["city", "pop"])
    ws.append(["lisbon", 545000])
    out = io.BytesIO()
    wb.save(out)

    text = await MultiFormatTextExtractor().extract(data=out.getvalue(), content_type=_XLSX_MIME)

    assert "Sheet1" in text
    assert "lisbon" in text
    assert "545000" in text


async def test_extract_docx_joins_paragraphs() -> None:
    document = Document()
    document.add_paragraph("hello from docx")
    document.add_paragraph("second paragraph")
    out = io.BytesIO()
    document.save(out)

    text = await MultiFormatTextExtractor().extract(data=out.getvalue(), content_type=_DOCX_MIME)

    assert "hello from docx" in text
    assert "second paragraph" in text


async def test_extract_pdf_reads_text() -> None:
    out = io.BytesIO()
    pdf = canvas.Canvas(out)
    pdf.drawString(100, 750, "cyberdyne attachment pdf")
    pdf.save()

    text = await MultiFormatTextExtractor().extract(
        data=out.getvalue(), content_type="application/pdf"
    )

    assert "cyberdyne attachment pdf" in text


async def test_extract_plain_text_decodes() -> None:
    text = await MultiFormatTextExtractor().extract(
        data=b"just plain text", content_type="text/plain"
    )
    assert text == "just plain text"


async def test_extract_unknown_type_returns_empty() -> None:
    text = await MultiFormatTextExtractor().extract(
        data=b"\x00\x01", content_type="application/octet-stream"
    )
    assert text == ""


async def test_extract_caps_output_length() -> None:
    big = ("x" * 50_000).encode("utf-8")
    text = await MultiFormatTextExtractor().extract(data=big, content_type="text/plain")
    assert len(text) == 20_000
